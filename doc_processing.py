from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx.shared import RGBColor

ADGM_CHECKLIST = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "Register of Members and Directors",
        "UBO Declaration Form"
    ]
}

def extract_text_from_docx(file_path):
    """Extract text from a .docx file."""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text])

def analyze_document(text, document_type):
    """Analyze the document for specific issues."""
    issues = []
    
    # Example checks
    if "jurisdiction" in text.lower() and "adgm" not in text:
        issues.append({
            "section": "Jurisdiction Clause",
            "issue": "Jurisdiction not specified as ADGM",
            "severity": "High",
            "suggestion": "Specify ADGM as the governing jurisdiction"
        })
    
    return issues

def identify_document_type(filename):
    """Identify the document type based on the filename."""
    filename = filename.lower()
    if "articles of association" in filename:
        return "Articles of Association"
    elif "memorandum" in filename:
        return "Memorandum of Association"
    elif "incorporation application form" in filename:
        return "Incorporation Application Form"
    elif "register of members and directors" in filename:
        return "Register of Members and Directors"
    elif "ubo declaration form" in filename:
        return "UBO Declaration Form"
    return "Unknown Document Type"

def process_documents(files):
    """Process a list of uploaded document files."""
    results = {}
    for file in files:
        text = extract_text_from_docx(file["name"])
        doc_type = identify_document_type(file["name"])
        issues = analyze_document(text, doc_type)
        results[file["name"]] = {
            "document_type": doc_type,
            "issues_found": issues
        }
    return results

def check_required_documents(uploaded_files):
    """Check if all required documents are present."""
    required_docs = ADGM_CHECKLIST["Company Incorporation"]
    uploaded_types = [identify_document_type(f["name"]) for f in uploaded_files]
    
    missing = [doc for doc in required_docs if doc not in uploaded_types]
    return missing

def detect_red_flags(doc):
    """Detect red flags in the document."""
    issues = []
    if "UAE Federal Courts" in doc:
        issues.append({
            "document": "General Document",
            "issue": "Incorrect jurisdiction referenced.",
            "severity": "High",
            "suggestion": "Update to ADGM Courts."
        })
    return issues

def add_comments_to_document(doc, comments):
    """Add comments to the document."""
    for para in doc.paragraphs:
        for comment in comments:
            if comment in para.text:
                run = para.add_run(f" [Comment: {comment}]")
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for comments

def save_marked_up_document(doc, output_path):
    """Save the marked-up document to the specified path."""
    doc.save(output_path)

def validate_documents(uploaded_files):
    """Validate uploaded documents and return analysis results."""
    results = process_documents(uploaded_files)
    missing_docs = check_required_documents(uploaded_files)
    
    # Compile final results
    complete = len(missing_docs) == 0
    return {
        "complete": complete,
        "jurisdiction_specified": all("ADGM" in text for text in results.values()),
        "issues": [issue for result in results.values() for issue in result["issues_found"]],
        "checklist_complete": complete
    }

if __name__ == "__main__":
    uploaded_files = [
        {"name": "Articles of Association.docx"},
        {"name": "Memorandum of Association.docx"},
        {"name": "Board Resolution.docx"},
        {"name": "Incorporation Application Form.docx"},
        {"name": "Register of Members and Directors.docx"}
    ]

    result = validate_documents(uploaded_files)
    print(result)
